"""多格式文档正文提取：docx / pdf / txt / md / 网页。"""
from __future__ import annotations

import io

import requests


def extract_text_from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    out = []
    for i, page in enumerate(reader.pages, 1):
        txt = page.extract_text() or ""
        out.append(f"\n[第{i}页]\n{txt}")
    return "\n".join(out)


def extract_text_from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    out = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "") if para.style else ""
        if style.startswith("Heading"):
            level = style.replace("Heading", "").strip() or "1"
            try:
                lv = int(level)
            except ValueError:
                lv = 1
            out.append(f"\n{'#' * min(lv, 6)} {text}\n")
        else:
            out.append(text)
    # 表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            out.append(" | ".join(cells))
    return "\n".join(out)


def extract_text_from_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore")


def extract_text_from_md(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore")


def extract_from_url(url: str) -> tuple[str, str]:
    """返回 (title, text)。优先 trafilatura，失败回退 BeautifulSoup。"""
    title, text = "", ""
    try:
        import trafilatura

        downloaded = trafilatura.fetch_url(url, timeout=20)
        if downloaded:
            text = trafilatura.extract(downloaded) or ""
    except Exception:
        text = ""
    if not text:
        try:
            resp = requests.get(url, timeout=20, headers={"User-Agent": "Mozilla/5.0"})
            resp.encoding = resp.apparent_encoding or "utf-8"
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(resp.text, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            title = soup.title.string if soup.title else ""
            text = soup.get_text(separator="\n")
        except Exception:
            text = ""
    return (title or url).strip(), text.strip()


_EXT_MAP = {
    ".pdf": ("pdf", extract_text_from_pdf),
    ".docx": ("file", extract_text_from_docx),
    ".txt": ("file", extract_text_from_txt),
    ".md": ("file", extract_text_from_md),
    ".markdown": ("file", extract_text_from_md),
}


def dispatch(source_type: str, data: bytes | None = None, url: str | None = None,
             filename: str = "") -> tuple[str, str]:
    """返回 (title, text)。source_type ∈ {file, url}。"""
    if source_type == "url":
        if not url:
            raise ValueError("url 类型必须提供 url")
        return extract_from_url(url)
    # file
    lowered = filename.lower()
    for ext, (_, fn) in _EXT_MAP.items():
        if lowered.endswith(ext):
            title = filename
            return title, fn(data)
    # 默认按纯文本
    return filename, extract_text_from_txt(data or b"")
